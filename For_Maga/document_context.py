"""
DocumentContextManager — загрузка пользовательских документов (PDF/DOCX/TXT/MD),
однократное извлечение текста и лёгкий локальный retrieval (TF-IDF + cosine
similarity, с fallback на пересечение ключевых слов, если scikit-learn не
установлен).

Ничего не отправляется во внешние embedding-сервисы — весь retrieval
локальный.
"""

from __future__ import annotations

import re
import threading
from dataclasses import dataclass, field
from pathlib import Path

CHUNK_SIZE = 1600
CHUNK_OVERLAP = 200

# Если суммарный объём документов укладывается в этот лимит - отправляем
# текст целиком, без retrieval.
FULL_CONTEXT_LIMIT = 30_000
DEFAULT_MAX_CHARS = 12_000

SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".txt", ".md")


@dataclass
class DocChunk:
    doc_id: str
    doc_name: str
    index: int
    text: str


@dataclass
class Document:
    doc_id: str
    path: str
    display_name: str
    size: int
    text: str = ""
    chunks: list[DocChunk] = field(default_factory=list)
    status: str = "pending"  # pending / ok / error
    error: str = ""


def _extract_pdf(path: str) -> str:
    from pypdf import PdfReader

    reader = PdfReader(path)
    parts = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text:
            parts.append(text)
    return "\n".join(parts).strip()


def _extract_docx(path: str) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(path)
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text for c in row.cells if c.text]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def _extract_text(path: str) -> str:
    data = Path(path).read_bytes()
    for enc in ("utf-8", "utf-8-sig", "cp1251"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _chunk_text(text: str, doc_id: str, doc_name: str) -> list[DocChunk]:
    chunks: list[DocChunk] = []
    if not text:
        return chunks
    start = 0
    idx = 0
    n = len(text)
    while start < n:
        end = min(start + CHUNK_SIZE, n)
        piece = text[start:end].strip()
        if piece:
            chunks.append(DocChunk(doc_id=doc_id, doc_name=doc_name, index=idx, text=piece))
            idx += 1
        if end >= n:
            break
        start = end - CHUNK_OVERLAP
    return chunks


def _keyword_overlap_scores(query: str, chunks: list[DocChunk]) -> list[float]:
    q_words = set(re.findall(r"\w+", query.lower()))
    if not q_words:
        return [0.0] * len(chunks)
    scores = []
    for c in chunks:
        c_words = set(re.findall(r"\w+", c.text.lower()))
        scores.append(float(len(q_words & c_words)))
    return scores


def _tfidf_scores(query: str, chunks: list[DocChunk]) -> list[float] | None:
    try:
        from sklearn.feature_extraction.text import TfidfVectorizer
        from sklearn.metrics.pairwise import cosine_similarity
    except ImportError:
        return None

    texts = [c.text for c in chunks] + [query]
    try:
        vectorizer = TfidfVectorizer()
        matrix = vectorizer.fit_transform(texts)
    except ValueError:
        return None
    query_vec = matrix[-1]
    doc_vecs = matrix[:-1]
    return list(cosine_similarity(query_vec, doc_vecs)[0])


def _retrieve(query: str, chunks: list[DocChunk], max_chars: int) -> str:
    scores = _tfidf_scores(query, chunks)
    if scores is None:
        scores = _keyword_overlap_scores(query, chunks)

    ranked = sorted(zip(scores, chunks), key=lambda x: x[0], reverse=True)
    out: list[str] = []
    used = 0
    for score, chunk in ranked:
        if score <= 0:
            continue
        piece = f"[{chunk.doc_name}]\n{chunk.text}"
        if out and used + len(piece) > max_chars:
            break
        out.append(piece)
        used += len(piece)
        if used >= max_chars:
            break
    return "\n\n---\n\n".join(out)


def _truncate_in_order(chunks: list[DocChunk], max_chars: int) -> str:
    out: list[str] = []
    used = 0
    for chunk in chunks:
        piece = f"[{chunk.doc_name}]\n{chunk.text}"
        if out and used + len(piece) > max_chars:
            break
        out.append(piece)
        used += len(piece)
        if used >= max_chars:
            break
    return "\n\n---\n\n".join(out)


class DocumentContextManager:
    """Хранит загруженные документы и отдаёт релевантный контекст по запросу."""

    def __init__(self) -> None:
        self._lock = threading.Lock()
        self._docs: dict[str, Document] = {}
        self._order: list[str] = []
        self._next_id = 1

    def add_files(self, paths: list[str]) -> list[Document]:
        """Регистрирует и извлекает текст из файлов. Блокирующий вызов —
        должен выполняться в background-потоке, не в Tkinter main thread."""
        added: list[Document] = []
        with self._lock:
            for p in paths:
                path = Path(p)
                doc_id = str(self._next_id)
                self._next_id += 1
                try:
                    size = path.stat().st_size
                except OSError:
                    size = 0
                doc = Document(doc_id=doc_id, path=str(path), display_name=path.name, size=size)
                self._docs[doc_id] = doc
                self._order.append(doc_id)
                added.append(doc)

        for doc in added:
            self._extract(doc)
        return added

    def _extract(self, doc: Document) -> None:
        ext = Path(doc.path).suffix.lower()
        try:
            if ext == ".pdf":
                text = _extract_pdf(doc.path)
            elif ext == ".docx":
                text = _extract_docx(doc.path)
            elif ext in (".txt", ".md"):
                text = _extract_text(doc.path)
            else:
                raise ValueError(f"неподдерживаемый формат: {ext or '?'}")
            doc.text = text
            doc.chunks = _chunk_text(text, doc.doc_id, doc.display_name)
            doc.status = "ok" if text else "empty"
        except Exception as e:  # noqa: BLE001
            doc.status = "error"
            doc.error = f"{type(e).__name__}: {e}"

    def remove_file(self, doc_id: str) -> None:
        with self._lock:
            self._docs.pop(doc_id, None)
            if doc_id in self._order:
                self._order.remove(doc_id)

    def clear(self) -> None:
        with self._lock:
            self._docs.clear()
            self._order.clear()

    def list_documents(self) -> list[Document]:
        with self._lock:
            return [self._docs[i] for i in self._order if i in self._docs]

    def has_documents(self) -> bool:
        with self._lock:
            return bool(self._order)

    def get_context(self, query: str, max_chars: int = DEFAULT_MAX_CHARS) -> str:
        """Возвращает релевантный контекст (или пусто, если документов нет).

        Небольшие наборы документов отправляются целиком. Для больших наборов
        выполняется локальный TF-IDF retrieval по `query`. Пустой `query`
        (например, до анализа скриншота) даёт компромиссный общий контекст
        без retrieval.
        """
        docs = [d for d in self.list_documents() if d.status == "ok" and d.text]
        if not docs:
            return ""

        total_len = sum(len(d.text) for d in docs)
        if total_len <= FULL_CONTEXT_LIMIT:
            joined = "\n\n".join(f"[{d.display_name}]\n{d.text}" for d in docs)
            return joined[:max_chars] if len(joined) > max_chars else joined

        all_chunks = [c for d in docs for c in d.chunks]
        if not all_chunks:
            return ""
        if not query or not query.strip():
            return _truncate_in_order(all_chunks, max_chars)
        return _retrieve(query, all_chunks, max_chars)
